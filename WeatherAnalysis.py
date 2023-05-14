#!/usr/bin/env python
# coding: utf-8

# In[1]:

import pandas as pd
import numpy as np
import os

import streamlit as st
from ladybug.epw import EPW
import pathlib
from plotly.graph_objects import Figure
from typing import List, Tuple

from ladybug.datacollection import HourlyContinuousCollection
from ladybug.epw import EPWFields
from ladybug.color import Colorset, Color
from ladybug.legend import LegendParameters
from ladybug.hourlyplot import HourlyPlot
from ladybug.monthlychart import MonthlyChart
from ladybug.analysisperiod import AnalysisPeriod

st.set_page_config(page_title='EPW Vizualiser Toolkit', layout='wide')

st.subheader('**Download the .epw File from the map**')


with st.container():
    
    st.components.v1.iframe("https://www.ladybug.tools/epwmap/", height = 800)


with st.sidebar:
    st.header('__EPW Visualiser Toolkit__')
    st.markdown('_Developed by **Amir Tabadkani**, \nPh.D. Computational Design Lead, Sustainability_')
    # st.write('Source codes: Ladybug Tools Core SDK Documentation')
#st.sidebar.image('https://www.ceros.com/wp-content/uploads/2019/04/Stantec_Logo.png',use_column_width='auto',output_format='PNG')

# Loading colorsets used in Legend Parameters
#---------------------------------------------------------------------------------
colorsets = {
    'original': Colorset.original(),
    'nuanced': Colorset.nuanced(),
    'annual_comfort': Colorset.annual_comfort(),
    'benefit': Colorset.benefit(),
    'benefit_harm': Colorset.benefit_harm(),
    'black_to_white': Colorset.black_to_white(),
    'blue_green_red': Colorset.blue_green_red(),
    'cloud_cover': Colorset.cloud_cover(),
    'cold_sensation': Colorset.cold_sensation(),
    'ecotect': Colorset.ecotect(),
    'energy_balance': Colorset.energy_balance(),
    'energy_balance_storage': Colorset.energy_balance_storage(),
    'glare_study': Colorset.glare_study(),
    'harm': Colorset.harm(),
    'heat_sensation': Colorset.heat_sensation(),
    'multi_colored': Colorset.multi_colored(),
    'multicolored_2': Colorset.multicolored_2(),
    'multicolored_3': Colorset.multicolored_3(),
    'openstudio_palette': Colorset.openstudio_palette(),
    'peak_load_balance': Colorset.peak_load_balance(),
    'shade_benefit': Colorset.shade_benefit(),
    'shade_benefit_harm': Colorset.shade_benefit_harm(),
    'shade_harm': Colorset.shade_harm(),
    'shadow_study': Colorset.shadow_study(),
    'therm': Colorset.therm(),
    'thermal_comfort': Colorset.thermal_comfort(),
    'view_study': Colorset.view_study()
}

# A function to derive the color code when selected in Streamlit
#------------------------------------------------------------------------------

@st.cache_resource(ttl=2,show_spinner=False)
def get_colors(switch: bool, global_colorset: str) -> List[Color]:
    """Get switched colorset if requested.
    Args:
        switch: Boolean to switch colorset.
        global_colorset: Global colorset to use.
    Returns:
        List of colors.
    """

    if switch:
        colors = list(colorsets[global_colorset])
        colors.reverse()
    else:
        colors = colorsets[global_colorset]
    return colors

# Define a function to extract the epw variable from the class
#------------------------------------------------------------------------------
@st.cache_resource(ttl=2,show_spinner=False)
def get_fields() -> dict:
    # A dictionary of EPW variable name to its corresponding field number
    return {EPWFields._fields[i]['name'].name: i for i in range(6, 34)}

# Uploading EPW file
#------------------------------------------------------------------------------

with st.sidebar:
    # epw file #####################################################################
    with st.expander('Upload EPW file'):
        epw_data = st.file_uploader('', type='epw')
        if epw_data:
            epw_file = pathlib.Path(f'./data/{epw_data.name}')
            epw_file.parent.mkdir(parents=True, exist_ok=True)
            epw_file.write_bytes(epw_data.read())
        else:
            epw_file = './assets/sample.epw'

        global_epw = EPW(epw_file)
    
    st.markdown('---')

# Global Colorset - Choose the heatmap color
#------------------------------------------------------------------------------

with st.sidebar:
    with st.expander('Global colorset'):
        global_colorset = st.selectbox('', list(colorsets.keys()))
        
# Hourly Plots
#------------------------------------------------------------------------------
st.markdown('---')
st.write("""
# Periodic Weather Data Analysis
         
***
""")

with st.sidebar:
    # A dictionary of EPW variable name to its corresponding field number
    
    fields = get_fields()
    with st.expander('Periodic analysis'):
        hourly_selected = st.selectbox('Which variable to plot?',options=fields.keys())
        _wea_data = global_epw.import_data_by_field(fields[hourly_selected])
        var_unit = _wea_data.header._unit
                  
        data_plot_radio = st.radio('How to plot the data?', ['Hourly Plot','Mean Daily Plot', 'Line Plot'], index =0, key='data_plot')
        
        if (data_plot_radio == 'Hourly Plot'):
            
            data_final = _wea_data
            
            hourly_data_st_month = st.number_input(
            'Start month', min_value=1, max_value=12, value=1, key='hourly_data_st_month')
            hourly_data_end_month = st.number_input(
                'End month', min_value=1, max_value=12, value=12, key='hourly_data_end_month')
    
            hourly_data_st_day = st.number_input(
                'Start day', min_value=1, max_value=31, value=1, key='hourly_data_st_day')
            hourly_data_end_day = st.number_input(
                'End day', min_value=1, max_value=31, value=31, key='hourly_data_end_day')
    
            hourly_data_st_hour = st.number_input(
                'Start hour', min_value=0, max_value=23, value=0, key='hourly_data_st_hour')
            hourly_data_end_hour = st.number_input(
                'End hour', min_value=0, max_value=23, value=23, key='hourly_data_end_hour')
            
        elif (data_plot_radio == 'Mean Daily Plot'):
            
            data_final = _wea_data
            
            hourly_data_st_month = None
            hourly_data_end_month = None
            hourly_data_st_day = None
            hourly_data_end_day = None
            hourly_data_st_hour = None
            hourly_data_end_hour = None            
            
        elif (data_plot_radio == 'Line Plot'): 
            
            data_final = _wea_data
            
            hourly_data_st_month = None
            hourly_data_end_month = None
            hourly_data_st_day = None
            hourly_data_end_day = None
            hourly_data_st_hour = None
            hourly_data_end_hour = None
            

@st.cache_resource(ttl=2,show_spinner=False)

def get_hourly_data_figure(data_type:str,
        _wea_data: HourlyContinuousCollection, global_colorset: str,st_month: int, st_day: int, st_hour: int, end_month: int,
        end_day: int, end_hour: int) -> Figure:
    """Create heatmap from hourly data.
    Args:
        data: HourlyContinuousCollection object.
        global_colorset: A string representing the name of a Colorset.
        conditional_statement: A string representing a conditional statement.
        min: A string representing the lower bound of the data range.
        max: A string representing the upper bound of the data range.
        st_month: start month.
        st_day: start day.
        st_hour: start hour.
        end_month: end month.
        end_day: end day.
        end_hour: end hour.
    Returns:
        A plotly figure.
    """
    lb_lp = LegendParameters(colors=colorsets[global_colorset])
    
    lb_ap = AnalysisPeriod(st_month, st_day, st_hour, end_month, end_day, end_hour)
    
    
    colors = colorsets[global_colorset] 
    
    if data_type == 'Hourly Plot':
        hourly_data = _wea_data.filter_by_analysis_period(lb_ap)
        hourly_plot = HourlyPlot(hourly_data, legend_parameters=lb_lp)
        
        return hourly_plot.plot(title=str(_wea_data.header.data_type), show_title=True)
    
    elif data_type == 'Mean Daily Plot':
        return _wea_data.diurnal_average_chart(
            title=_wea_data.header.data_type.name, show_title=True,color=colors[-1])
    
    elif data_type == 'Line Plot':
        return _wea_data.line_chart(title=_wea_data.header.data_type.name,show_title=True,color=colors[-1])
    
Hourly_figure = get_hourly_data_figure(data_plot_radio,data_final,global_colorset, hourly_data_st_month, hourly_data_st_day,
                hourly_data_st_hour, hourly_data_end_month, hourly_data_end_day,
                hourly_data_end_hour)

st.header(f'{global_epw.location.city}, {global_epw.location.country}')
st.plotly_chart(Hourly_figure, use_container_width=True)

#Saving images

Hourly_figure_image = get_hourly_data_figure('Hourly Plot',data_final,global_colorset, hourly_data_st_month, hourly_data_st_day,
                hourly_data_st_hour, hourly_data_end_month, hourly_data_end_day,
                hourly_data_end_hour)
Daily_figure_image = get_hourly_data_figure('Mean Daily Plot',data_final,global_colorset, hourly_data_st_month, hourly_data_st_day,
                hourly_data_st_hour, hourly_data_end_month, hourly_data_end_day,
                hourly_data_end_hour)
Line_figure_image = get_hourly_data_figure('Line Plot',data_final,global_colorset, hourly_data_st_month, hourly_data_st_day,
                hourly_data_st_hour, hourly_data_end_month, hourly_data_end_day,
                hourly_data_end_hour)

Hourly_figure_image.write_image("hourly_data.png")
Daily_figure_image.write_image("Daily_data.png")
Line_figure_image.write_image("Line_data.png")


#statistics
from statistics import mean
ave_val = round(mean(_wea_data._values),2)


# CONDITIONAL HOURLY PLOTS
#------------------------------------------------------------------------------

with st.sidebar: 
    with st.expander('Conditional statement'):
        fields = get_fields()
        st.markdown(':red[**Min/Max Thresholds**]')
        
        min_value = _wea_data.bounds[0]
        max_value = _wea_data.bounds[1]
      
        threshold_min = st.slider('Minimum {}'.format(hourly_selected), min_value, max_value, value = min_value, step=None)
        threshold_max = st.slider('Maximum {}'.format(hourly_selected), min_value, max_value, value = max_value, step=None)
        
@st.cache_resource(ttl=2,show_spinner=False)

def get_hourly_data_figure_conditional(_hourly_data: HourlyContinuousCollection, global_colorset: str,st_month: int, st_day: int, st_hour: int, end_month: int,
        end_day: int, end_hour: int) -> Figure:
    """Create heatmap from hourly data.
    Args:
        data: HourlyContinuousCollection object.
        global_colorset: A string representing the name of a Colorset.
        conditional_statement: A string representing a conditional statement.
        min: A string representing the lower bound of the data range.
        max: A string representing the upper bound of the data range.
        st_month: start month.
        st_day: start day.
        st_hour: start hour.
        end_month: end month.
        end_day: end day.
        end_hour: end hour.
    Returns:
        A plotly figure.
    """
    lb_lp = LegendParameters(colors=colorsets[global_colorset])
    
    lb_ap = AnalysisPeriod(st_month, st_day, st_hour, end_month, end_day, end_hour)
    _hourly_data = _hourly_data.filter_by_analysis_period(lb_ap)
    
           
    hourly_plot = HourlyPlot(_hourly_data, legend_parameters=lb_lp)

    return hourly_plot.plot(title=str(_hourly_data.header.data_type), show_title=True)

st.subheader('_Applied Thresholds_')
st.markdown('Please choose the thresholds from the min/max sliders on the left to plot the filtered data below:')

with st.container():
    
    filtered_hourly_data= _wea_data.filter_by_analysis_period(AnalysisPeriod(hourly_data_st_month,hourly_data_st_day,hourly_data_st_hour,hourly_data_end_month,hourly_data_end_day,hourly_data_end_hour))
    
    data_work_hours = filtered_hourly_data.filter_by_conditional_statement('a>={} and a<={}'.format(threshold_min,threshold_max))
    
    Hourly_conditional_figure = get_hourly_data_figure_conditional(data_work_hours,global_colorset, hourly_data_st_month, hourly_data_st_day,
                    hourly_data_st_hour, hourly_data_end_month, hourly_data_end_day,
                    hourly_data_end_hour)
    st.plotly_chart(Hourly_conditional_figure, use_container_width=True)
    
    
    met_num_hours = len(data_work_hours)
    unmet_num_hours = len(filtered_hourly_data)-len(data_work_hours)
    
    col1,col2 = st.columns(2)
    with col1:
        st.metric(':blue[**Met hours for the selected {} thresholds:**]'.format(hourly_selected), value = met_num_hours)
    with col2:
        st.metric(':red[**Unmet hours for the selected {} thresholds:**]'.format(hourly_selected), value = unmet_num_hours)

Hourly_conditional_figure.write_image("conditional_hourly_data.png")
   
st.markdown('---')

#Pyschometric Chart
#------------------------------------------------------------------------------

import ladybug.psychrometrics
import ladybug_comfort
from ladybug.psychchart import PsychrometricChart
from ladybug_charts.utils import Strategy
from ladybug_comfort.chart.polygonpmv import PolygonPMV

st.write("""
# Psychrometric Chart
         
***
""")


with st.sidebar:
    
    with st.expander('Psychrometric chart'):
    
        fields = get_fields()
        

        psy_radio = st.radio('',['Load Hourly Data', 'Extracting Psychrometrics', 'Calculate PMV/PPD'],index = 0, key='psy_radio')
        
        if psy_radio == 'Load Hourly Data':
            psy_selected = st.selectbox('Select an environmental variable', options=fields.keys())
            psy_data = global_epw.import_data_by_field(fields[psy_selected])
            psy_draw_polygons = st.checkbox('Draw comfort polygons')
            psy_clo_value = st.number_input('Clothing Level',value=0.7)
            psy_met_value = st.number_input('Metabloic Rate',value=1.1)
            psy_air = st.number_input('Air Velocity (m/s)' ,value = 0.1)
            psy_strategy_options = ['Comfort', 'Evaporative Cooling',
                                    'Mass + Night Ventilation', 'Occupant use of fans',
                                    'Capture internal heat', 'Passive solar heating', 'All']
            psy_selected_strategy = st.selectbox(
                'Select a passive strategy', options=psy_strategy_options)
        elif psy_radio == 'Extracting Psychrometrics':
            psy_selected_strategy = None
            psy_draw_polygons = None
            psy_data = None
            psy_clo_value = None
            psy_met_value = None
            psy_air = None
            psy_db = st.number_input('Insert DBT value',min_value =-20, max_value = 50, value = 24)
            psy_rh = st.number_input('Insert RH value',min_value =0, max_value = 100, value = 45)
        
        elif psy_radio == 'Calculate PMV/PPD':
            psy_selected = None
            psy_selected_strategy = None
            psy_draw_polygons = None
            psy_data = None
            psy_db = st.number_input('DBT',min_value =-20, max_value = 50, value = 24)
            psy_rh = st.number_input('RH',min_value =0, max_value = 100, value = 45)
            psy_mrt = st.number_input('MRT',min_value =0, max_value = 100, value = psy_db)
            psy_air = st.number_input('Air Velocity (m/s)' ,value = 0.1)
            psy_clo_value = st.number_input('Clothing Level',value=0.7)
            psy_met_value = st.number_input('Metabloic Rate',value=1.1)
            
            
@st.cache_resource(ttl=2,show_spinner=False)

def get_psy_chart_figure(_epw: EPW, global_colorset: str, selected_strategy: str,
                         load_data: str, draw_polygons: bool,
                         _data: HourlyContinuousCollection) -> Tuple[Figure, HourlyContinuousCollection, Tuple]:
    """Create psychrometric chart figure.
    Args:
        epw: An EPW object.
        global_colorset: A string representing the name of a Colorset.
        selected_strategy: A string representing the name of a psychrometric strategy.
        load_data: A boolean to indicate whether to load the data.
        draw_polygons: A boolean to indicate whether to draw the polygons.
        data: Hourly data to load on psychrometric chart.
    Returns:
        A plotly figure.
    """

    lb_lp = LegendParameters(colors=colorsets[global_colorset])
    lb_psy = PsychrometricChart(_epw.dry_bulb_temperature,
                                _epw.relative_humidity, legend_parameters=lb_lp)
    
    if selected_strategy == 'All':
        strategies = [Strategy.comfort, Strategy.evaporative_cooling,
                      Strategy.mas_night_ventilation, Strategy.occupant_use_of_fans,
                      Strategy.capture_internal_heat, Strategy.passive_solar_heating]
    elif selected_strategy == 'Comfort':
        strategies = [Strategy.comfort]
    elif selected_strategy == 'Evaporative Cooling':
        strategies = [Strategy.evaporative_cooling]
    elif selected_strategy == 'Mass + Night Ventilation':
        strategies = [Strategy.mas_night_ventilation]
    elif selected_strategy == 'Occupant use of fans':
        strategies = [Strategy.occupant_use_of_fans]
    elif selected_strategy == 'Capture internal heat':
        strategies = [Strategy.capture_internal_heat]
    elif selected_strategy == 'Passive solar heating':
        strategies = [Strategy.passive_solar_heating]
        
    pmv_param = ladybug_comfort.parameter.pmv.PMVParameter(10,humid_ratio_upper = 1, humid_ratio_lower=0)
    
    if load_data == 'Load Hourly Data':
        
        pmv = PolygonPMV(lb_psy,met_rate=[psy_met_value],clo_value=[psy_clo_value], air_speed = [psy_air], comfort_parameter = pmv_param )
        
        #Extracting values for the report
        internal_heat_pc = round(sum(pmv.evaluate_polygon(pmv.internal_heat_polygon(),0.01)),2)/100
        Fan_pc = round(sum(pmv.evaluate_polygon(pmv.fan_use_polygon(),0.01)),2)/100
        nightflush_pc = round(sum(pmv.evaluate_polygon(pmv.night_flush_polygon(),0.01)),2)/100
        evaluate_passive_solar = pmv.evaluate_passive_solar(global_epw.global_horizontal_radiation,50,8,12.8)
        passivesolar_pc = round(sum(pmv.evaluate_polygon(pmv.passive_solar_polygon(evaluate_passive_solar[1]),0.01)),2)/100
        evap_clg_pc = round(sum(pmv.evaluate_polygon(pmv.evaporative_cooling_polygon(),0.01)),2)/100
        
        comfort_value_pc = round((pmv.merged_comfort_data.total/8760)*100,2)
        
        strategies_percentages = [internal_heat_pc,Fan_pc,nightflush_pc,passivesolar_pc,evap_clg_pc,comfort_value_pc]
        
        if draw_polygons:
            
           
            figure = lb_psy.plot(data=_data, polygon_pmv=pmv,
                                 strategies=strategies,title='PSYCHROMETRIC CHART', show_title=True, solar_data =global_epw.global_horizontal_radiation )
            
        else:
            figure = lb_psy.plot(data=_data, show_title=True)
            
        return figure, strategies_percentages, None
    
    elif load_data == 'Extracting Psychrometrics':
        
        
        
        lb_psy_ext = PsychrometricChart(psy_db,psy_rh)
        figure = lb_psy_ext.plot(title='PSYCHROMETRIC CHART', show_title=True)
        
        dew_pt = round(ladybug.psychrometrics.dew_point_from_db_rh(psy_db, psy_rh),2)
        hr_pt = round(ladybug.psychrometrics.humid_ratio_from_db_rh(psy_db, psy_rh),2)
        wb_pt = round(ladybug.psychrometrics.wet_bulb_from_db_rh(psy_db, psy_rh),2)
        ent_pt = round(ladybug.psychrometrics.enthalpy_from_db_hr(psy_db, hr_pt),2)
        psy_db_kelvin = psy_db + 274.15
        sat_p = round(ladybug.psychrometrics.saturated_vapor_pressure(psy_db_kelvin),2)
        
        col1,col2,col3,col4,col5,col6,col7 = st.columns(7)
        
        with col1:
            
            st.metric('Dry Bulb Temperature (°C)', value = psy_db)
            
        with col2:
            
            st.metric('Relative Humidity (%)', value = psy_rh)
            
        with col3:
            
            st.metric('Dew Point Temperature (°C)', value = dew_pt)
            
        with col4:
            
            st.metric('Humidty Ratio (kg_H₂O kg_Air⁻¹)', value = hr_pt)
            
        with col5:
            
            st.metric('Wet Bulb Temperature (°C)', value = wb_pt)
            
        with col6:
            
            st.metric('Enthalpy (J kg⁻¹)', value = ent_pt)
            
        with col7:
            
            st.metric('Saturated Vapour Pressure (Pa)', value = sat_p)
        
        return figure, None,  None

    else:
        
        pmv = PolygonPMV(lb_psy,met_rate=[psy_met_value],clo_value=[psy_clo_value], air_speed = [psy_air], comfort_parameter = pmv_param )

        lb_psy_ext = PsychrometricChart(psy_db,psy_rh)
        figure = lb_psy_ext.plot(polygon_pmv=pmv, title='PSYCHROMETRIC CHART', show_title=True)
        

        PMV_cal = ladybug_comfort.pmv.fanger_pmv(psy_db, psy_mrt , psy_air , psy_rh, psy_met_value, psy_clo_value)
        
        col1,col2,col3,col4 = st.columns([2,1,1,2])
        
        with col1:
            st.markdown('')
            
        with col2:
            st.metric('**:orange[PMV Fanger Value]**', value = round(PMV_cal[0],2))
        
        with col3:
            st.metric('**:orange[PPD Fanger Value (%)]**',value = round(PMV_cal[1],2) )
            
        with col4:
            st.markdown('')
        
        return figure, None, PMV_cal

@st.cache_resource(ttl=2,show_spinner=False)

def get_figure_config(title: str) -> dict:
    """Set figure config so that a figure can be downloaded as SVG."""

    return {
        'toImageButtonOptions': {
            'format': 'svg',  # one of png, svg, jpeg, webp
            'filename': title,
            'height': 700,
            'width': 700,
            'scale': 1  # Multiply title/legend/axis/canvas sizes by this factor
        }
    }

with st.container():
    
    st.markdown('The following Psychrometric chart contains THREE functions here:')
    st.markdown(
        '(1) Load Hourly Data: A psychrometric chart plots multiple data '
        'points, that represent the air conditions at a specific time, on the chart and overlays an area '
        'that identifies the “comfort zone.”  The comfort zone is defined as the range within occupants are '
        'satisfied with the surrounding thermal conditions. After plotting the air conditions and overlaying '
        'the comfort zone, it becomes possible to see how passive design strategies can extend the comfort '
        'zone.')
    st.markdown(
        '(2) Extracting psychrometrics: Calculating the psychrometrics based on the given dry bulb temperature'
        ' and relative humidity')
    st.markdown(
        '(3) Calculate PMV/PPD: Calculating thermal comfort indices including Predicted Mean Vote (PMV) and Percentage of People Dissatisfied (PPD)' 
        ' based on the given necessary inputs for an indoor environmental condition.')
    
    psy_chart_figure, strategies_percentages, PMV_cal  = get_psy_chart_figure(
        global_epw, global_colorset, psy_selected_strategy, psy_radio,
        psy_draw_polygons,psy_data)
    
    
    st.plotly_chart(psy_chart_figure, use_container_width=True, config=get_figure_config(f'Psychrometric_chart_{global_epw.location.city}'))


#Saving image
psy_chart_figure.write_image("psy_main.png")

# WINDROSE
#------------------------------------------------------------------------------
from ladybug.windrose import WindRose

st.markdown('---')
st.write("""
# Wind Rose
""")
st.markdown('---')

with st.sidebar:
    
    with st.expander('Windrose'):
    
        windrose_st_month = st.number_input(
            'Start month', min_value=1, max_value=12, value=1, key='windrose_st_month')
        windrose_end_month = st.number_input(
            'End month', min_value=1, max_value=12, value=12, key='windrose_end_month')

        windrose_st_day = st.number_input(
            'Start day', min_value=1, max_value=31, value=1, key='windrose_st_day')
        windrose_end_day = st.number_input(
            'End day', min_value=1, max_value=31, value=31, key='windrose_end_day')

        windrose_st_hour = st.number_input(
            'Start hour', min_value=0, max_value=23, value=0, key='windrose_st_hour')
        windrose_end_hour = st.number_input(
            'End hour', min_value=0, max_value=23, value=23, key='windrose_end_hour')
    
   
@st.cache_resource(ttl=2,show_spinner=False)

def get_windrose_figure(st_month: int, st_day: int, st_hour: int, end_month: int,
                        end_day: int, end_hour: int, _epw, global_colorset) -> Figure:
    
    """Create windrose figure.
    Args:
        st_month: A number representing the start month.
        st_day: A number representing the start day.
        st_hour: A number representing the start hour.
        end_month: A number representing the end month.
        end_day: A number representing the end day.
        end_hour: A number representing the end hour.
        epw: An EPW object.
        global_colorset: A string representing the name of a Colorset.
    Returns:
        A plotly figure.
    """
            
    lb_ap = AnalysisPeriod(st_month, st_day, st_hour, end_month, end_day, end_hour)
    wind_dir = _epw.wind_direction.filter_by_analysis_period(lb_ap)
    wind_spd = _epw.wind_speed.filter_by_analysis_period(lb_ap)
    
    lb_lp = LegendParameters(colors=colorsets[global_colorset])
    
    lb_wind_rose = WindRose(wind_dir, wind_spd)
    lb_wind_rose.legend_parameters = lb_lp
    
    
    return lb_wind_rose.plot(title='Windrose',show_title=True)

@st.cache_resource(ttl=2,show_spinner=False)

def get_windrose_figure_temp(st_month: int, st_day: int, st_hour: int, end_month: int,
                    end_day: int, end_hour: int, _epw, global_colorset) -> Figure:
    
    """Create windrose figure.
    Args:
        st_month: A number representing the start month.
        st_day: A number representing the start day.
        st_hour: A number representing the start hour.
        end_month: A number representing the end month.
        end_day: A number representing the end day.
        end_hour: A number representing the end hour.
        epw: An EPW object.
        global_colorset: A string representing the name of a Colorset.
    Returns:
        A plotly figure.
    """
    
    lb_ap = AnalysisPeriod(st_month, st_day, st_hour, end_month, end_day, end_hour)        
    
    fields = get_fields()
        
    windrose_data = global_epw.import_data_by_field(fields['Dry Bulb Temperature'])
    
    wind_dir = _epw.wind_direction.filter_by_analysis_period(lb_ap)
    windrose_data_ = windrose_data.filter_by_analysis_period(lb_ap)
    
    lb_lp = LegendParameters(colors=colorsets[global_colorset])
    
    lb_windrose_temp = WindRose(wind_dir,windrose_data_)
    lb_windrose_temp.legend_parameters = lb_lp
    
    return lb_windrose_temp.plot(title='Wind Direction vs. Dry Bulb Temperature', show_title=True)
 
@st.cache_resource(ttl=2,show_spinner=False)

def get_windrose_figure_dir_rad(st_month: int, st_day: int, st_hour: int, end_month: int,
                    end_day: int, end_hour: int, _epw, global_colorset) -> Figure:
    
    """Create windrose figure.
    Args:
        st_month: A number representing the start month.
        st_day: A number representing the start day.
        st_hour: A number representing the start hour.
        end_month: A number representing the end month.
        end_day: A number representing the end day.
        end_hour: A number representing the end hour.
        epw: An EPW object.
        global_colorset: A string representing the name of a Colorset.
    Returns:
        A plotly figure.
    """
    
    lb_ap = AnalysisPeriod(st_month, st_day, st_hour, end_month, end_day, end_hour)        
    
    fields = get_fields()
        
    windrose_data = global_epw.import_data_by_field(fields['Direct Normal Radiation'])
    
    wind_dir = _epw.wind_direction.filter_by_analysis_period(lb_ap)
    windrose_data_ = windrose_data.filter_by_analysis_period(lb_ap)
    
    lb_lp = LegendParameters(colors=colorsets[global_colorset])
    
    lb_windrose_dir = WindRose(wind_dir,windrose_data_)
    lb_windrose_dir.legend_parameters = lb_lp
    
    return lb_windrose_dir.plot(title='Wind Direction vs. Direct Normal Radiation', show_title=True)

@st.cache_resource(ttl=2,show_spinner=False)

def get_windrose_figure_diff_rad(st_month: int, st_day: int, st_hour: int, end_month: int,
                    end_day: int, end_hour: int, _epw, global_colorset) -> Figure:
    
    """Create windrose figure.
    Args:
        st_month: A number representing the start month.
        st_day: A number representing the start day.
        st_hour: A number representing the start hour.
        end_month: A number representing the end month.
        end_day: A number representing the end day.
        end_hour: A number representing the end hour.
        epw: An EPW object.
        global_colorset: A string representing the name of a Colorset.
    Returns:
        A plotly figure.
    """
    
    lb_ap = AnalysisPeriod(st_month, st_day, st_hour, end_month, end_day, end_hour)        
    
    fields = get_fields()
        
    windrose_data = global_epw.import_data_by_field(fields['Diffuse Horizontal Radiation'])
    
    wind_dir = _epw.wind_direction.filter_by_analysis_period(lb_ap)
    windrose_data_ = windrose_data.filter_by_analysis_period(lb_ap)
    
    lb_lp = LegendParameters(colors=colorsets[global_colorset])
    
    lb_windrose_diff = WindRose(wind_dir,windrose_data_)
    lb_windrose_diff.legend_parameters = lb_lp
    
    return lb_windrose_diff.plot(title='Wind Direction vs. Diffuse Horizontal Radiation', show_title=True)


with st.container():
    
    st.markdown('Generate a wind rose to summarise the occurrence of winds at a location, showing their strength, direction and frequency, for the selected period and environmental parameter!')
    
    col1,col2,col3,col4 = st.columns(4)
    with col1:
        
        windrose_figure = get_windrose_figure(windrose_st_month, windrose_st_day, windrose_st_hour, windrose_end_month,
                                              windrose_end_day, windrose_end_hour, global_epw, global_colorset)
    
        st.plotly_chart(windrose_figure, use_container_width=True,
                        config=get_figure_config(f'Windrose_{global_epw.location.city}'))
    with col2:
        windrose_figure_temp = get_windrose_figure_temp(windrose_st_month, windrose_st_day, windrose_st_hour, windrose_end_month,
                                              windrose_end_day, windrose_end_hour, global_epw, global_colorset)
    
        st.plotly_chart(windrose_figure_temp, use_container_width=True,
                        config=get_figure_config(f'Windrose_{global_epw.location.city}'))
    
    with col3:
        windrose_figure_dir = get_windrose_figure_dir_rad(windrose_st_month, windrose_st_day, windrose_st_hour, windrose_end_month,
                                              windrose_end_day, windrose_end_hour, global_epw, global_colorset)
    
        st.plotly_chart(windrose_figure_dir, use_container_width=True,
                        config=get_figure_config(f'Windrose_{global_epw.location.city}'))
        
    with col4:
        windrose_figure_diff = get_windrose_figure_diff_rad(windrose_st_month, windrose_st_day, windrose_st_hour, windrose_end_month,
                                              windrose_end_day, windrose_end_hour, global_epw, global_colorset)
    
        st.plotly_chart(windrose_figure_diff, use_container_width=True,
                        config=get_figure_config(f'Windrose_{global_epw.location.city}'))
        
    st.markdown('---')

#Saving image
windrose_figure.write_image("windrose.png")
windrose_figure_temp.write_image("windrose-temp.png")
windrose_figure_dir.write_image("windrose-dir.png")
windrose_figure_diff.write_image("windrose-diff.png")

#SUNPATH
#-----------------------------------------------------------------------------
from ladybug.sunpath import Sunpath

st.write("""
# Sunpath Diagram
         
***
""")
with st.sidebar:
    
    with st.expander('Sunpath'):
    
        sunpath_radio = st.radio(
            '', ['from epw location', 'with epw data'],
            index=0, key='sunpath_'
        )
    
        if sunpath_radio == 'from epw location':
            sunpath_switch = st.checkbox('Switch colors', key='sunpath_switch',
                                         help='Reverse the colorset')
            sunpath_data = None
    
        else:
            sunpath_selected = st.selectbox(
                'Select an environmental variable', options=fields.keys(), key='sunpath')
            sunpath_data = global_epw._get_data_by_field(fields[sunpath_selected])
            sunpath_switch = None
            
@st.cache_resource(ttl=2,show_spinner=False)

def get_sunpath_figure(sunpath_type: str, global_colorset: str, _epw: EPW = None,
                       switch: bool = False,
                       _data: HourlyContinuousCollection = None, ) -> Figure:
    """Create sunpath figure.
    Args:
        sunpath_type: A string representing the type of sunpath to be plotted.
        global_colorset: A string representing the name of a Colorset.
        epw: An EPW object.
        switch: A boolean to indicate whether to reverse the colorset.
        data: Hourly data to load on sunpath.
    Returns:
        A plotly figure.
    """
    if sunpath_type == 'from epw location':
        lb_sunpath = Sunpath.from_location(_epw.location)
        colors = get_colors(switch, global_colorset)
        return lb_sunpath.plot(title='Sunpath Diagram',colorset=colors,show_title=True)
    else:
        lb_sunpath = Sunpath.from_location(_epw.location)
        colors = colorsets[global_colorset]
        return lb_sunpath.plot(title=str(_data.header.data_type),colorset=colors, data=_data, show_title=True)

with st.container():

    st.markdown('Generate a sunpath using EPW location. Additionally, you can'
                ' also load one of the environmental variables from the EPW file'
                ' on the sunpath.'
                )
    
    sunpath_figure = get_sunpath_figure(sunpath_radio, global_colorset, global_epw, sunpath_switch, sunpath_data)
    
    st.plotly_chart(sunpath_figure, use_container_width=True,
                    config=get_figure_config(
                        f'Sunpath_{global_epw.location.city}'))


#Saving images
sunpath_figure.write_image("Sunpath.png")


#Degree Days
#------------------------------------------------------------------------------


from ladybug.datatype.temperaturetime import HeatingDegreeTime, CoolingDegreeTime
from ladybug_comfort.degreetime import heating_degree_time, cooling_degree_time

with st.sidebar:
    with st.expander('Degree days'):
    
          
        degree_days_heat_base = st.number_input('Base heating temperature',
                                                value=18)

    
        degree_days_cool_base = st.number_input('Base cooling temperature',
                                                value=23)
        
        dd_st_hour = st.number_input(
            'Start hour', min_value=0, max_value=23, value=0, key='dd_st_hour')
        dd_end_hour = st.number_input(
            'End hour', min_value=0, max_value=23, value=23, key='dd_end_hour')
        
@st.cache_resource(ttl=2,show_spinner=False)

def get_degree_days_figure(_st_hour: int, _end_hour: int,
    _heat_base_: int, _cool_base_: int,
    global_colorset: str) -> Tuple[Figure,HourlyContinuousCollection,HourlyContinuousCollection]:
    """Create HDD and CDD figure.
    Args:
        dbt: A HourlyContinuousCollection object.
        _heat_base_: A number representing the heat base temperature.
        _cool_base_: A number representing the cool base temperature.
        global_colorset: A string representing the name of a Colorset.
    Returns:
        A tuple of three items:
        -   A plotly figure.
        -   Heating degree days as a HourlyContinuousCollection.
        -   Cooling degree days as a HourlyContinuousCollection.
    """
    lb_ap = AnalysisPeriod(1, 1, _st_hour, 12, 31, _end_hour) 
    
    filtered_dd = global_epw.dry_bulb_temperature.filter_by_analysis_period(lb_ap)
    
    hourly_heat = filtered_dd.compute_function_aligned(
        heating_degree_time, [filtered_dd, _heat_base_],
        HeatingDegreeTime(), 'degC-hours')
    hourly_heat.convert_to_unit('degC-days')

    hourly_cool = filtered_dd.compute_function_aligned(
        cooling_degree_time, [filtered_dd, _cool_base_],
        CoolingDegreeTime(), 'degC-hours')
    hourly_cool.convert_to_unit('degC-days')

    colors = colorsets[global_colorset]

    lb_lp = LegendParameters(colors=colors)
    monthly_chart = MonthlyChart([hourly_cool.total_monthly(),
                                  hourly_heat.total_monthly()], legend_parameters=lb_lp)

    return monthly_chart.plot(title='Degree Days'), hourly_heat, hourly_cool

with st.container():
    st.markdown('---')
    st.header('Degree Days')
    st.markdown('---')
    st.markdown('**Degree days** is another way of combining time and temperature, but it has different implications for heating and cooling than does design temperature.')
                
    st.markdown('**Heating degree days (HDD)** – This is the number that tells you' 
                'how long a location stays below a special temperature called the base'
                'temperature. I find it easiest to start with degree hours.'
                'For example, the most commonly used base temperature for heating is 18°C.'
                'So if the temperature at your house is 12° F for one hour, you just'
                'accumulated 6 degree hours. If the temperature is 15°C for the next hour'
                ',you’ve got another 3 degree hours and 9 degree hours total.'
                'To find the number of degree days, you divide it by 24,'
                ' so you’ve got one degree day in this example.'
                'You can do that for every hour of the year to find the total and then divide by 24.'
                ' Or you can use the average temperature for each day to get degree days directly.')
    st.markdown('**Cooling degree days (CDD)** – Same principle as for heating degree days but usually'
                'with a different base temperature which is here set as 23°C by default.') 
                

    degree_days_figure, hourly_heat, hourly_cool = get_degree_days_figure(dd_st_hour,dd_end_hour,
        degree_days_heat_base,degree_days_cool_base,global_colorset)

    st.plotly_chart(degree_days_figure, use_container_width=True,
                    config=get_figure_config(
                        f'Degree days_{global_epw.location.city}'))
    col1, col2  = st.columns(2)
    with col1:
        st.metric(':blue[**TOTOAL COOLING DEGREE HOURS**]', value = round(hourly_cool.total))
    with col2:
        st.metric(':red[**TOTAL HEATING DEGREE HOURS**]', value = round(hourly_heat.total))


#Saving image
degree_days_figure.write_image("CDD_HDD.png")

#Distributed DBT Plot
#------------------------------------------------------------------------------
import plotly.express as px


    
with st.sidebar:
    
    with st.expander('Temperature Range Settings'):
    
        min_val_bin = st.number_input("Minimum Value", min_value = -20, max_value = 0, value = 0)
        max_val_bin = st.number_input("Maximum Value", min_value = 20, max_value = 60, value = 40)
        steps = st.slider("Number of Steps", min_value = 1, max_value = 5, value = 2)

with st.container():
    
    st.markdown('---')
    st.header('Distributed Temperature Plot')
    st.markdown('---')
    
    dbt = global_epw.dry_bulb_temperature
     
    db_df = pd.DataFrame(list(dbt.values), columns = ['Dry Bulb Temperature'])
    
    bins = list(np.arange(min_val_bin,max_val_bin,steps))
    
    @st.cache_data(ttl=2,show_spinner=False)

    def get_ranges():
        ranges = []
        for i in range(0,(len(bins)-1)):
            x = str(f'{bins[i]}-{bins[i]+steps}°C')
            ranges.append(x)
        return ranges
    
    ranges_str = get_ranges()
    
    db_df['Temperature Range'] = pd.cut(db_df['Dry Bulb Temperature'], bins, labels = ranges_str)
    db_df = db_df.groupby('Temperature Range').count()
    
    temp_bins_plot = px.bar(db_df)
    temp_bins_plot = temp_bins_plot.update_traces(textfont_size=12, textangle=0, textposition="outside", cliponaxis=False)
    
    st.plotly_chart(temp_bins_plot, use_container_width=True)


Maximum_bin = db_df.index[(db_df["Dry Bulb Temperature"].argmax())]
Minimum_bin = db_df.index[(db_df["Dry Bulb Temperature"].argmin())]


#Saving image
temp_bins_plot.write_image('temp-bins.png')

#Generate the REPORT in WORD
#------------------------------------------------------------------------------

from docx import Document
from docx.shared import Inches


export_as_docs = st.button("Export Report (.docx)")

document = Document()

x = 2
h_res = x
w_res = 3*x

document.add_heading(f'Weather Analysis for {global_epw.location.city} ', level= 0)

p1 = document.add_paragraph('This report outlines the Weather Analysis for ')
p1.add_run(f'{global_epw.location.city}').bold = True
p1.add_run(' located in ')
p1.add_run(f'{global_epw.location.country}').bold = True
p1.add_run(' in the following sections: ')

document.add_paragraph('Periodic Weather Data Analysis', style='List Number')
p2 = document.add_paragraph('The following graphs illustrate the ')
p2.add_run(f'{hourly_selected}').bold = True
p2.add_run(f' variable for the selected period in which the average {hourly_selected} is {ave_val}{var_unit} for the year,'
           f' and ranges from {min_value}{var_unit} up to {max_value}{var_unit} (Figure 1).'
           f' Figure 2 and Figure 3 show the mean daily values and {hourly_selected} variation per day.')

document.add_picture('hourly_data.png', width=Inches(w_res), height= Inches(h_res))
document.add_paragraph(f'Figure 1. Hourly {hourly_selected} Analysis', style='Caption')

document.add_picture('Daily_data.png', width=Inches(w_res), height= Inches(h_res))
document.add_paragraph(f'Figure 2. Mean Daily {hourly_selected} Analysis', style='Caption')

document.add_picture('Line_data.png', width=Inches(w_res), height= Inches(h_res))
document.add_paragraph(f'Figure 3. Average/Range {hourly_selected} Analysis', style='Caption')

document.add_paragraph(f'Conditional Hourly Data: {hourly_selected}', style='List Number')
p3 = document.add_paragraph('Below graph plots the given thresholds between')
p3.add_run(f' {threshold_min}{var_unit}').bold= True
p3.add_run(' and ')
p3.add_run(f'{threshold_max}{var_unit}').bold = True 
p3.add_run(f' for {hourly_selected}. Additionally, the number of met hours within the given range is {met_num_hours}hrs and unmet hours is {unmet_num_hours}hrs.')

document.add_picture('conditional_hourly_data.png', width=Inches(w_res), height= Inches(h_res))
document.add_paragraph(f'Figure 4. Hourly Conditional Analysis of {hourly_selected}', style='Caption')

document.add_paragraph('Psychrometric Chart', style='List Number')
document.add_paragraph('A psychrometric chart can be used in two different ways.'
                            ' The first is done by plotting multiple data points, that represent'
                            ' the air conditions at a specific time, on the chart. Then, overlaying'
                            ' an area that identifies the “comfort zone.” The comfort zone is defined '
                            'as the range within occupants are satisfied with the surrounding thermal '
                            'conditions. After plotting the air conditions and overlaying the comfort '
                            'zone, it becomes possible to see how passive design strategies can extend '
                            'the comfort zone.')

document.add_picture('psy_main.png', width=Inches(w_res), height= Inches(h_res*2))
document.add_paragraph('Figure 5. Psychrometric Chart Analysis', style='Caption')

document.add_paragraph(f'The chart above depicts the hourly distribution of the outdoor air condition throughout the year associated with hourly relative humidty and humidity ratio.'
                            f' In particular, having a Clothing Level of {psy_clo_value} and Metabloic Rate of {psy_met_value}'
                            ' resulted in the following statistics:')
if psy_radio == 'Load Hourly Data':
    
    document.add_paragraph(f'{strategies_percentages[5]}% of the time is in comfortable range without the need of any passive/active strategies.', style = 'List Bullet')
    document.add_paragraph(f'{strategies_percentages[0]}% of the time, internal heat gains can improve the environmental condition to meet the comfortable range.', style = 'List Bullet')
    document.add_paragraph(f'{strategies_percentages[1]}% of the time, occupants can use fans (assuming 1m/s of air velocity around the occupant) to reduce the temperature and improve thermal comfort.', style = 'List Bullet')
    document.add_paragraph(f'{strategies_percentages[2]}% of the time, thermal massing and night ventilation allow the possibility to maintain the heat during the day and release it during the night to reduce the temperature passively.', style = 'List Bullet')
    document.add_paragraph(f'{strategies_percentages[3]}% of the time, passive solar heat gains can increase the temperature as an additional potential. Noting that it assumes a 50W/m2 outdoor solar flux (W/m2) that is needed to raise the temperature of a theoretical building by 1 degree Celsius and can maintain its temperature for 8 hours.', style = 'List Bullet')
    document.add_paragraph(f'{strategies_percentages[4]}% of the time, evaporative cooling strategy has the opportunity to cool down the building while increasing the relative humidity.', style = 'List Bullet')
else:
    with st.container():
        
        st.error('In order to have the analysis in line with the generated report, we suggest to select'+' Load Hourly Data Option '+'for Psychrometric chart analysis!', icon="❌")
    
document.add_paragraph('Wind Rose Diagrams', style='List Number')
document.add_paragraph('Wind roses are graphical charts that characterize the speed and direction'
                       ' of winds at a location. Presented in a circular format, the length of each'
                       ' "spoke" around the circle indicates the amount of time that the wind blows'
                       ' from a particular direction. Colors along the spokes indicate categories'
                       ' of wind speed. Additionally, the other diagrams plot the correlation between the wind direction and'
                       ' Dry Bulb Temperature, Direct/Diffuse Solar Radiation which is helpful to understand'
                       ' the possibilities of utilizing natural ventilation and shading strategies.')

document.add_picture('windrose.png', width=Inches(4.2), height= Inches(3))
document.add_paragraph('Figure 6. Wind Direction vs. Dry Bulb Temperature', style='Caption')
document.add_picture('windrose-temp.png', width=Inches(4.2), height= Inches(3))
document.add_paragraph('Figure 7. Wind Rose vs. Wind Direction', style='Caption')
document.add_picture('windrose-dir.png', width=Inches(4.2), height= Inches(3))
document.add_paragraph('Figure 8. Wind Rose vs. Direct Solar Radiation', style='Caption')
document.add_picture('windrose-diff.png', width=Inches(4.2), height= Inches(3))
document.add_paragraph('Figure 9. Wind Rose vs. Diffuse Solar Radiation', style='Caption')

document.add_paragraph('Sunpath Diagram', style='List Number')
document.add_paragraph('Solar charts and sun path diagrams have been devised as visual'
                       'aids so that the solar position can be easily and quickly established'
                       'for any hour in any day, to provide a valuable tool for designers and planners.'
                       'A separate sun path diagram is required for each latitude.'
                       ' Additionally, plotting one of the environmental variables on the sunpath diagram'
                       ' can help to understand its variation across the sun position.')

document.add_picture('Sunpath.png', width=Inches(4.2), height= Inches(3))
document.add_paragraph('Figure 10. Sunpath Diagram', style='Caption')

document.add_paragraph('Cooling/Heating Degree Days/Hours', style='List Number')
document.add_paragraph('Degree days is another way of combining time and temperature, but it has different implications for heating and cooling than does design temperature.'
                       ' Heating degree days (HDD) – This is the number that tells you' 
                'how long a location stays below a special temperature called the base'
                'temperature. I find it easiest to start with degree hours.'
                'For example, the most commonly used base temperature for heating is 18°C.'
                'So if the temperature at your house is 12°C for one hour, you just'
                ' accumulated 6 degree hours. If the temperature is 15°C for the next hour'
                ', you’ve got another 3 degree hours and 9 degree hours total.'
                'To find the number of degree days, you divide it by 24,'
                ' so you’ve got one degree day in this example.'
                'You can do that for every hour of the year to find the total and then divide by 24.'
                ' Or you can use the average temperature for each day to get degree days directly.'
                'Same principle as for heating degree days applies to Cooling degree days (CDD) but '
                'with a different base temperature which is typically set as 24°C by default.') 
document.add_paragraph(f'Following the selected base temperatures, total HEATING HOURS and COOLING HOURS are respetively {round(hourly_heat.total,0)} and {round(hourly_heat.total,0)} in {global_epw.location.city}')
document.add_picture('CDD_HDD.png', width=Inches(w_res), height= Inches(h_res*1.5))
document.add_paragraph('Figure 11. Cooling/Heating Degree Days', style='Caption')


document.add_paragraph('Distributed Temperature Plot', style='List Number')
p4 = document.add_paragraph('To better understand the outdoor temperature levels, the figure below divides the dry'
                       f' bulb temperature into sequential bins from the minimum value of {min_val_bin}°C up to the maximum value of {max_val_bin}°C'
                       f' where the highest and lowest intensities are')
p4.add_run(f' {Maximum_bin} and ').bold= True   
p4.add_run(f'{Minimum_bin}').bold= True          
p4.add_run(' respectively')              
                       

document.add_picture('temp-bins.png', width=Inches(w_res), height= Inches(h_res*1.5))
document.add_paragraph('Figure 12. Temperature Ranges', style='Caption')



if export_as_docs:
    filepath = pathlib.Path.home()  
    document.save(pathlib.Path(filepath,"Downloads", f'WeatherAnalysis-{global_epw.location.city}.docx'))

pathlib.Path()
st.markdown('Please note that the generated report will take your inputs as the basis of the weather analysis. Therefore, make sure you have selected the right values/thresholds and proper environmental variables given in the control panel based on your design needs.')
st.markdown('**The REPORT is in your DOWNLOADS folder now, ENJOY READING!**')

